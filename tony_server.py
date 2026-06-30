"""
TONY — local assistant backend (v3, agentic)
Adds the agent core: tool calling, a planner loop, and multi-agent delegation,
on top of v2's document memory (RAG), web search, and Groq/Ollama brains.

Tools available to the agent:
  - web_search        (free, no key)
  - search_documents  (your indexed PDFs/Word/code via ChromaDB)
  - run_python        (executes Python locally — calculations, file tasks)
  - read_webpage      (fetch & read any URL)
  - get_datetime      (current date/time)
  - delegate          (spin up a focused sub-agent: research / code / plan)
  - browser_*         (real browser control — needs Playwright, see below)
  - calendar_*/email_*(Google — activates after you complete OAuth, see below)

Run:
    python tony_server.py        (with Ollama running)  →  http://localhost:8000

Install (one-time):
    pip install flask flask-cors requests ddgs chromadb sentence-transformers pypdf python-docx beautifulsoup4

Browser control (optional, for tool 3):
    pip install playwright
    playwright install chromium

Fast brain (recommended for agents — tool calling is sharper on a bigger model):
    set GROQ_API_KEY  (free key at https://console.groq.com/keys), then pick groq:... in the app.
"""

import os, sys, json, glob, subprocess, time, threading
import requests
from flask import Flask, request, Response, send_from_directory, g, make_response
from flask_cors import CORS
from urllib.parse import urlparse

try:
    from ddgs import DDGS
except ImportError:
    from duckduckgo_search import DDGS

OLLAMA = "http://localhost:11434"
PORT = int(os.environ.get("PORT", "8000"))
GROQ_KEY = os.environ.get("GROQ_API_KEY", "").strip()
OPENAI_KEY = os.environ.get("OPENAI_API_KEY", "").strip()
GEMINI_KEY = os.environ.get("GEMINI_API_KEY", "").strip() or os.environ.get("GOOGLE_API_KEY", "").strip()
ANTHROPIC_KEY = os.environ.get("ANTHROPIC_API_KEY", "").strip()
XAI_KEY = os.environ.get("XAI_API_KEY", "").strip() or os.environ.get("GROK_API_KEY", "").strip()
SPOTIFY_ID = os.environ.get("SPOTIFY_CLIENT_ID", "").strip()
SPOTIFY_SECRET = os.environ.get("SPOTIFY_CLIENT_SECRET", "").strip()
SPOTIFY_REDIRECT = os.environ.get("SPOTIFY_REDIRECT", "http://127.0.0.1:%s/spotify/callback" % os.environ.get("PORT", "8000"))
# Gmail (read-only) — set GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET, then visit /google/login once
GOOGLE_ID = os.environ.get("GOOGLE_CLIENT_ID", "").strip()
GOOGLE_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "").strip()
GOOGLE_REDIRECT = os.environ.get("GOOGLE_REDIRECT", "http://127.0.0.1:%s/google/callback" % os.environ.get("PORT", "8000"))
GOOGLE_SCOPES = "https://www.googleapis.com/auth/gmail.readonly"
GOOGLE_TOK = os.path.join(os.path.dirname(os.path.abspath(__file__)), "google_token.json")
# CareerOS project folder — Tony watches it for changes and folds updates into the briefing
CAREEROS_DIR = os.environ.get("TONY_CAREEROS", r"K:\Projects\CareerOS")

# OpenAI-compatible providers: prefix -> (base_url, key)
COMPAT = {
    "groq":   ("https://api.groq.com/openai/v1", lambda: GROQ_KEY),
    "openai": ("https://api.openai.com/v1", lambda: OPENAI_KEY),
    "gemini": ("https://generativelanguage.googleapis.com/v1beta/openai", lambda: GEMINI_KEY),
    "grok":   ("https://api.x.ai/v1", lambda: XAI_KEY),
}

def split_model(model):
    for pfx in ("groq", "openai", "gemini", "grok", "claude"):
        if model.startswith(pfx + ":"):
            return pfx, model[len(pfx) + 1:]
    return "ollama", model
DB_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tony_db")
MAX_STEPS = 6  # agent loop guard

def _api_error(r):
    """Turn a failed provider response into a short human message."""
    try:
        j = r.json()
    except Exception:
        return f"HTTP {r.status_code}: {(r.text or '')[:200]}"
    if isinstance(j, dict) and j.get("error"):
        e = j["error"]
        return e.get("message") if isinstance(e, dict) else str(e)
    return f"HTTP {r.status_code}"

def free_chain():
    """Free / effectively-unlimited brains, in preferred order, that have a key or run locally."""
    chain = []
    if GROQ_KEY:   chain += ["groq:llama-3.3-70b-versatile", "groq:llama-3.1-8b-instant"]
    if GEMINI_KEY: chain += ["gemini:gemini-2.0-flash"]
    chain += ["phi3:mini"]          # local Ollama — offline, unlimited, last resort
    return chain

def build_chain(model):
    """Requested model first, then automatic free fallbacks (keeps you chatting if one errors)."""
    chain = [model]
    for m in free_chain():
        if m not in chain:
            chain.append(m)
    return chain

app = Flask(__name__, static_folder=".")

# ===========================================================================
#  Security — proportionate for a personal, local-first assistant
#  (NOT enterprise public-service hardening; see notes in chat)
# ===========================================================================
#  HOST   : localhost by default — unreachable from other devices.
#           Set TONY_LAN=1 (or TONY_HOST=0.0.0.0) ONLY when you want phone access.
#  TOKEN  : if TONY_TOKEN is set, other devices must present it once (?token=...).
#           Your own machine (127.0.0.1) never needs it.
#  Origin : cross-site POSTs are refused, so a random website you visit can't
#           quietly call your local server's run_python/agent (real local risk).
#  Rate   : simple in-memory cap so nothing can hammer the heavy endpoints.
_cloud = bool(os.environ.get("RENDER") or os.environ.get("DYNO") or
              os.environ.get("FLY_APP_NAME") or os.environ.get("K_SERVICE"))
HOST  = os.environ.get("TONY_HOST", "") or ("0.0.0.0" if (os.environ.get("TONY_LAN") in ("1","true","yes") or _cloud) else "127.0.0.1")
TOKEN = os.environ.get("TONY_TOKEN", "").strip()
RATE  = int(os.environ.get("TONY_RATE", "240"))          # requests per IP per minute
WRITE_PATHS = ("/chat", "/agent", "/ingest", "/speak", "/jobs/refresh", "/feed/run", "/save",
               "/spotify/play", "/spotify/pause", "/spotify/next", "/spotify/prev")
SEC_LOG = os.path.join(os.path.dirname(os.path.abspath(__file__)), "security.log")
_hits = {}                                               # ip -> [timestamps]

# CORS: same-origin app needs none; allow only localhost dev origins, not "*"
CORS(app, origins=[f"http://localhost:{int(os.environ.get('PORT','8000'))}",
                   f"http://127.0.0.1:{int(os.environ.get('PORT','8000'))}"])

def _seclog(reason, ip):
    try:
        with open(SEC_LOG, "a", encoding="utf-8") as f:
            f.write(f"{time.strftime('%Y-%m-%d %H:%M:%S')}  {ip:<15}  {reason}  "
                    f"{request.method} {request.path}  ua={request.headers.get('User-Agent','')[:80]}\n")
    except Exception:
        pass

def _local(ip):
    return ip in ("127.0.0.1", "::1", "localhost")

@app.before_request
def _guard():
    ip = (request.headers.get("X-Forwarded-For", request.remote_addr) or "").split(",")[0].strip()
    # 1) cross-site write guard — block a foreign website from POSTing to us
    if request.method == "POST":
        origin = request.headers.get("Origin", "")
        if origin and urlparse(origin).netloc and urlparse(origin).netloc != request.host:
            _seclog("blocked cross-origin POST origin=" + origin, ip)
            return Response('{"error":"cross-origin request blocked"}', status=403, mimetype="application/json")
    # 2) rate limit the heavy endpoints
    if request.path in WRITE_PATHS:
        now = time.time(); win = _hits.setdefault(ip, [])
        win[:] = [t for t in win if now - t < 60]
        if len(win) >= RATE:
            _seclog("rate-limited", ip)
            return Response('{"error":"too many requests"}', status=429, mimetype="application/json")
        win.append(now)
    # 3) optional access token for remote devices (your own machine is exempt)
    if TOKEN and not _local(ip):
        supplied = (request.cookies.get("tony_token") or request.headers.get("X-Tony-Token")
                    or request.args.get("token") or "")
        if supplied != TOKEN:
            _seclog("denied (bad/missing token)", ip)
            return Response("Tony is locked. Open this device's link with ?token=YOUR_TOKEN once.",
                            status=401, mimetype="text/plain")
        if request.args.get("token") == TOKEN:
            g.set_token = True       # first valid visit → remember on this device

@app.after_request
def _harden(resp):
    if getattr(g, "set_token", False):
        resp.set_cookie("tony_token", TOKEN, max_age=60*60*24*30, samesite="Strict", httponly=True)
    resp.headers["X-Content-Type-Options"] = "nosniff"
    resp.headers["Referrer-Policy"] = "no-referrer"
    return resp

# ===========================================================================
#  Document memory (RAG)
# ===========================================================================
_collection = None
def get_collection():
    global _collection
    if _collection is not None:
        return _collection
    try:
        import chromadb
        from chromadb.utils import embedding_functions
        client = chromadb.PersistentClient(path=DB_DIR)
        ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
        _collection = client.get_or_create_collection("tony_docs", embedding_function=ef)
    except Exception as e:
        print("  [RAG disabled]", e); _collection = None
    return _collection

def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            from pypdf import PdfReader
            return "\n".join((p.extract_text() or "") for p in PdfReader(path).pages)
        if ext == ".docx":
            import docx
            return "\n".join(p.text for p in docx.Document(path).paragraphs)
        if ext in (".txt",".md",".markdown",".py",".js",".ts",".json",".csv",".html",
                   ".htm",".css",".xml",".yaml",".yml",".log",".java",".c",".cpp",".h",".sql",".sh"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        print("  [skip]", path, e)
    return ""

def chunk(t, size=900, overlap=150):
    t = " ".join(t.split()); out, i = [], 0
    while i < len(t):
        out.append(t[i:i+size]); i += size - overlap
    return out

def retrieve(query, k=4):
    col = get_collection()
    if col is None or col.count() == 0:
        return "No documents indexed yet."
    try:
        res = col.query(query_texts=[query], n_results=k)
        docs = res.get("documents", [[]])[0]; metas = res.get("metadatas", [[]])[0]
        if not docs:
            return "Nothing relevant found in your documents."
        return "\n\n".join(f"[{m.get('name','file')}] {d[:1000]}" for d, m in zip(docs, metas))
    except Exception as e:
        return "Document search error: " + str(e)

# ===========================================================================
#  Markdown memory vault — plain .md files. NO embedder, NO database, no
#  model download. This is Tony's everyday memory: notes you ask him to keep
#  and the reports his routines generate are written as markdown and recalled
#  by keyword search. Works even when the ChromaDB embedder isn't installed.
# ===========================================================================
import re as _re
VAULT_DIR = os.environ.get("TONY_VAULT") or os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "tony_vault")
os.makedirs(VAULT_DIR, exist_ok=True)

_VAULT_STOP = set((
    "a an the is are was were be been being do does did to of in on at for and or but "
    "if then with as by from this that these those it its i you my your me we our he she "
    "they them his her their what which who whom whose how when where why can could should "
    "would will just about into over under not no yes so out up down all any have has had"
).split())

def _vault_slug(title):
    s = _re.sub(r"[^a-z0-9]+", "-", (title or "note").lower()).strip("-")
    return (s or "note")[:60]

def _vault_words(text):
    return [w for w in _re.findall(r"[a-z0-9]+", (text or "").lower())
            if len(w) > 2 and w not in _VAULT_STOP]

def vault_save(title, content, kind="note", unique=False):
    """Write a markdown note. If a note with the same title already exists and
    unique is False, append a new dated section so memory accumulates rather
    than overwriting. Reports use unique=True to keep every copy."""
    os.makedirs(VAULT_DIR, exist_ok=True)
    slug = _vault_slug(title)
    if unique:
        slug = slug + "-" + time.strftime("%Y%m%d-%H%M%S")
    path = os.path.join(VAULT_DIR, slug + ".md")
    stamp = time.strftime("%Y-%m-%d %H:%M")
    content = (content or "").strip()
    if os.path.exists(path) and not unique:
        with open(path, "a", encoding="utf-8") as f:
            f.write("\n\n## %s\n\n%s\n" % (stamp, content))
    else:
        with open(path, "w", encoding="utf-8") as f:
            f.write("---\ntitle: %s\nkind: %s\ncreated: %s\n---\n\n# %s\n\n%s\n"
                    % (title, kind, stamp, title, content))
    return {"ok": True, "slug": slug, "path": path, "name": slug + ".md"}

def _vault_files():
    return sorted(glob.glob(os.path.join(VAULT_DIR, "*.md")),
                  key=os.path.getmtime, reverse=True)

def vault_list(limit=80):
    out = []
    for p in _vault_files()[:limit]:
        try:
            head = open(p, encoding="utf-8").read(500)
        except Exception:
            head = ""
        m = _re.search(r"title:\s*(.+)", head)
        title = m.group(1).strip() if m else os.path.basename(p)[:-3]
        out.append({"name": os.path.basename(p), "title": title,
                    "modified": time.strftime("%Y-%m-%d %H:%M", time.localtime(os.path.getmtime(p))),
                    "size": os.path.getsize(p)})
    return out

def vault_search(query, max_hits=4):
    """Keyword search across markdown notes — no embeddings required. Scores
    each file by how often the query's words appear (filename/title matches
    weighted higher), then returns the single most relevant paragraph from each
    of the top files. Returns (text_block, hits[])."""
    qwords = _vault_words(query)
    if not qwords:
        return "", []
    scored = []
    for p in _vault_files():
        try:
            text = open(p, encoding="utf-8").read()
        except Exception:
            continue
        low = text.lower()
        score = sum(low.count(w) for w in qwords)
        base = os.path.basename(p).lower()
        score += 3 * sum(1 for w in qwords if w in base)
        if score > 0:
            scored.append((score, p, text))
    scored.sort(key=lambda x: x[0], reverse=True)
    blocks, hits = [], []
    for _, p, text in scored[:max_hits]:
        paras = [pp.strip() for pp in _re.split(r"\n\s*\n", text)
                 if pp.strip() and not pp.strip().startswith("---")]
        best = max(paras, key=lambda pp: sum(pp.lower().count(w) for w in qwords)) if paras else text[:600]
        snippet = best[:600]
        name = os.path.basename(p)[:-3]
        blocks.append("[%s] %s" % (name, snippet))
        hits.append({"name": name, "snippet": snippet})
    return "\n\n".join(blocks), hits

# ===========================================================================
#  Skills — SKILL.md folders. Each skills/<name>/SKILL.md describes a
#  capability and its triggers. The matching skill (and only that one) is
#  loaded into context for a reply, so Tony stays focused.
# ===========================================================================
SKILLS_DIR = os.environ.get("TONY_SKILLS") or os.path.join(
    os.path.dirname(os.path.abspath(__file__)), "skills")

def _parse_skill(path):
    raw = open(path, encoding="utf-8").read()
    name = os.path.basename(os.path.dirname(path))
    triggers, body = [], raw
    m = _re.match(r"\s*---(.*?)---(.*)", raw, _re.S)
    if m:
        fm, body = m.group(1), m.group(2).strip()
        nm = _re.search(r"name:\s*(.+)", fm)
        if nm: name = nm.group(1).strip()
        tg = _re.search(r"triggers:\s*(.+)", fm)
        if tg: triggers = [t.strip().lower() for t in tg.group(1).split(",") if t.strip()]
    if not triggers: triggers = [name.lower()]
    return {"name": name, "triggers": triggers, "body": body, "path": path}

def load_skills():
    out = []
    for md in sorted(glob.glob(os.path.join(SKILLS_DIR, "*", "SKILL.md"))):
        try: out.append(_parse_skill(md))
        except Exception as e: print("  [skill]", e)
    return out

SKILLS = load_skills()
if SKILLS:
    print("  Skills loaded:", ", ".join(s["name"] for s in SKILLS))

def match_skill(text):
    low = (text or "").lower()
    for s in SKILLS:
        if any(t in low for t in s["triggers"]):
            return s
    return None

# ===========================================================================
#  Tools
# ===========================================================================
def t_web_search(query=""):
    try:
        out = []
        with DDGS() as d:
            for r in d.text(query, max_results=5):
                out.append(f"- {r.get('title','')}: {r.get('body','')[:200]} ({r.get('href','')})")
        return "\n".join(out) or "No results."
    except Exception as e:
        return "Search failed: " + str(e)

def t_search_documents(query=""):
    return retrieve(query)

def t_run_python(code=""):
    try:
        r = subprocess.run([sys.executable, "-c", code], capture_output=True, text=True, timeout=20)
        out = (r.stdout or "") + (("\n[stderr] " + r.stderr) if r.stderr else "")
        return out.strip()[:3000] or "(ran, no output)"
    except subprocess.TimeoutExpired:
        return "Timed out (20s limit)."
    except Exception as e:
        return "Python error: " + str(e)

def t_read_webpage(url=""):
    try:
        h = {"User-Agent": "Mozilla/5.0"}
        html = requests.get(url, headers=h, timeout=15).text
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            for s in soup(["script", "style", "nav", "footer"]):
                s.extract()
            text = " ".join(soup.get_text(" ").split())
        except Exception:
            import re
            text = " ".join(re.sub("<[^>]+>", " ", html).split())
        return text[:4000]
    except Exception as e:
        return "Couldn't read page: " + str(e)

def t_get_datetime(_=""):
    import datetime
    return datetime.datetime.now().strftime("%A %d %B %Y, %H:%M")

# ---- Playwright browser control (optional) ----
_browser = _page = None
def _ensure_browser():
    global _browser, _page
    if _page:
        return True
    from playwright.sync_api import sync_playwright
    pw = sync_playwright().start()
    _browser = pw.chromium.launch(headless=False)
    _page = _browser.new_page()
    return True

def t_browser(action="", url="", text=""):
    try:
        _ensure_browser()
        if action == "open":
            _page.goto(url, timeout=30000); return "Opened " + url + " — title: " + _page.title()
        if action == "read":
            return " ".join(_page.inner_text("body").split())[:4000]
        if action == "click":
            _page.click(f"text={text}", timeout=8000); return "Clicked: " + text
        if action == "type":
            _page.keyboard.type(text); return "Typed: " + text
        return "Unknown browser action."
    except ModuleNotFoundError:
        return "Browser control needs Playwright. Run: pip install playwright && playwright install chromium"
    except Exception as e:
        return "Browser error: " + str(e)

# ---- Google calendar / email (activate after OAuth) ----
def t_calendar(action="list", **kw):
    token = os.path.join(os.path.dirname(__file__), "google_token.json")
    if not os.path.exists(token):
        return ("Calendar not connected yet. One-time setup: enable Google Calendar API, download "
                "credentials.json into this folder, then run the google_setup.py helper to authorise. "
                "Until then I can't read your calendar.")
    try:
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
        import datetime
        creds = Credentials.from_authorized_user_file(token)
        svc = build("calendar", "v3", credentials=creds)
        now = datetime.datetime.utcnow().isoformat() + "Z"
        ev = svc.events().list(calendarId="primary", timeMin=now, maxResults=8,
                               singleEvents=True, orderBy="startTime").execute().get("items", [])
        if not ev:
            return "No upcoming events."
        return "\n".join(f"- {e['start'].get('dateTime', e['start'].get('date'))}: {e.get('summary','(no title)')}" for e in ev)
    except Exception as e:
        return "Calendar error: " + str(e)

def define_word(word):
    word = (word or "").strip().split()[0] if (word or "").strip() else ""
    if not word:
        return "Which word would you like me to define?"
    try:
        r = requests.get("https://api.dictionaryapi.dev/api/v2/entries/en/" + word, timeout=12)
        if r.status_code != 200:
            return "No dictionary entry for “%s”." % word
        j = r.json()[0]; head = word
        ph = j.get("phonetic") or next((p.get("text") for p in j.get("phonetics", []) if p.get("text")), "")
        if ph: head += "  " + ph
        out = [head]
        for m in j.get("meanings", [])[:3]:
            d = (m.get("definitions") or [{}])[0]
            out.append("(%s) %s" % (m.get("partOfSpeech", ""), d.get("definition", "")))
            if d.get("example"): out.append("   e.g. " + d["example"])
        return "\n".join(out)
    except Exception as e:
        return "Dictionary lookup failed: " + str(e)

TOOLS_SPEC = [
    {"type": "function", "function": {"name": "web_search", "description": "Search the live internet for current information.",
        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]}}},
    {"type": "function", "function": {"name": "search_documents", "description": "Search Bharadwaj's own indexed documents (PDFs, CVs, papers, code).",
        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]}}},
    {"type": "function", "function": {"name": "run_python", "description": "Execute Python code locally and return its output. Use for maths, data, file tasks.",
        "parameters": {"type": "object", "properties": {"code": {"type": "string"}}, "required": ["code"]}}},
    {"type": "function", "function": {"name": "read_webpage", "description": "Fetch and read the text of a specific URL.",
        "parameters": {"type": "object", "properties": {"url": {"type": "string"}}, "required": ["url"]}}},
    {"type": "function", "function": {"name": "get_datetime", "description": "Get the current date and time.",
        "parameters": {"type": "object", "properties": {}}}},
    {"type": "function", "function": {"name": "delegate", "description": "Hand a focused sub-task to a specialist sub-agent. role is one of: research, code, plan.",
        "parameters": {"type": "object", "properties": {"role": {"type": "string"}, "task": {"type": "string"}}, "required": ["role", "task"]}}},
    {"type": "function", "function": {"name": "browser", "description": "Control a real Chrome browser. action: open|read|click|type. Use url for open, text for click/type.",
        "parameters": {"type": "object", "properties": {"action": {"type": "string"}, "url": {"type": "string"}, "text": {"type": "string"}}, "required": ["action"]}}},
    {"type": "function", "function": {"name": "calendar", "description": "Read Bharadwaj's Google Calendar (upcoming events).",
        "parameters": {"type": "object", "properties": {"action": {"type": "string"}}}}},
    {"type": "function", "function": {"name": "play_music", "description": "Play a song or artist on the user's Spotify (Premium). query = what to play; leave empty to resume.",
        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}}}},
    {"type": "function", "function": {"name": "dictionary", "description": "Look up the precise dictionary definition and pronunciation of a single word.",
        "parameters": {"type": "object", "properties": {"word": {"type": "string"}}, "required": ["word"]}}},
    {"type": "function", "function": {"name": "save_file", "description": "Save text content to a downloadable file on the user's laptop. format: txt, md, docx, or pdf.",
        "parameters": {"type": "object", "properties": {"content": {"type": "string"}, "filename": {"type": "string"}, "format": {"type": "string"}}, "required": ["content"]}}},
    {"type": "function", "function": {"name": "check_email", "description": "Read recent job-related emails (recruiters, applications, interviews). Read-only, never sends.",
        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}}}},
    {"type": "function", "function": {"name": "note_save", "description": "Save something to long-term markdown memory (the vault). Use when Bharadwaj says remember this / make a note, or after producing something worth keeping.",
        "parameters": {"type": "object", "properties": {"title": {"type": "string"}, "content": {"type": "string"}}, "required": ["content"]}}},
    {"type": "function", "function": {"name": "note_search", "description": "Search Bharadwaj's saved markdown notes and past reports (the memory vault) by keyword. Use to recall earlier notes, decisions, or briefings.",
        "parameters": {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]}}},
]

def run_tool(name, args, allow_delegate=True):
    if name == "web_search":      return t_web_search(args.get("query", ""))
    if name == "search_documents":return t_search_documents(args.get("query", ""))
    if name == "run_python":      return t_run_python(args.get("code", ""))
    if name == "read_webpage":    return t_read_webpage(args.get("url", ""))
    if name == "get_datetime":    return t_get_datetime()
    if name == "browser":         return t_browser(args.get("action", ""), args.get("url", ""), args.get("text", ""))
    if name == "calendar":        return t_calendar(args.get("action", "list"))
    if name == "play_music":      return sp_play_text(args.get("query", ""))
    if name == "dictionary":      return define_word(args.get("word", ""))
    if name == "save_file":
        r = save_document(args.get("content", ""), args.get("filename", "tony_note"), args.get("format", "txt"))
        return ("Saved %s → %s" % (r["name"], r["url"]) + (" (" + r["note"] + ")" if r.get("note") else "")) if r.get("ok") else ("Couldn't save: " + r.get("error", ""))
    if name == "check_email":     return email_summary_text(args.get("query", ""))
    if name == "note_save":
        r = vault_save(args.get("title", "Note") or "Note", args.get("content", ""))
        return "Saved to your memory vault: " + r["name"]
    if name == "note_search":
        block, _ = vault_search(args.get("query", ""), max_hits=4)
        return block or "No matching notes in the vault yet."
    if name == "delegate":
        if not allow_delegate:    return "Delegation depth limit reached."
        return run_subagent(args.get("role", "research"), args.get("task", ""))
    return "Unknown tool: " + name

# ===========================================================================
#  LLM call (normalised across Groq + Ollama) with tool support
# ===========================================================================
def llm(model, messages, tools=None):
    """Returns (text, tool_calls) where tool_calls = [{'name':..,'args':{}}]."""
    prov, mid = split_model(model)
    if prov in COMPAT:
        base, keyfn = COMPAT[prov]
        body = {"model": mid, "messages": messages}
        if tools: body["tools"] = tools; body["tool_choice"] = "auto"
        r = requests.post(base + "/chat/completions",
                          headers={"Authorization": "Bearer " + keyfn()}, json=body, timeout=120)
        j = r.json() if r.content else {}
        if r.status_code != 200 or not isinstance(j, dict) or "choices" not in j:
            raise RuntimeError(_api_error(r))
        m = j["choices"][0]["message"]
        calls = [{"name": c["function"]["name"], "args": json.loads(c["function"]["arguments"] or "{}")}
                 for c in m.get("tool_calls", []) or []]
        return m.get("content") or "", calls
    if prov == "claude":
        sys_txt = " ".join(x["content"] for x in messages if x.get("role") == "system" and isinstance(x.get("content"), str))
        conv = [{"role": x["role"], "content": x["content"]} for x in messages
                if x.get("role") in ("user", "assistant") and isinstance(x.get("content"), str)]
        r = requests.post("https://api.anthropic.com/v1/messages",
                          headers={"x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"},
                          json={"model": mid, "system": sys_txt, "messages": conv, "max_tokens": 1024}, timeout=120)
        txt = "".join(b.get("text", "") for b in r.json().get("content", []))
        return txt, []
    # ollama
    body = {"model": mid, "messages": messages, "stream": False}
    if tools: body["tools"] = tools
    r = requests.post(OLLAMA + "/api/chat", json=body, timeout=300)
    m = r.json().get("message", {})
    calls = [{"name": c["function"]["name"],
              "args": c["function"].get("arguments", {}) if isinstance(c["function"].get("arguments"), dict)
              else json.loads(c["function"].get("arguments") or "{}")}
             for c in m.get("tool_calls", []) or []]
    return m.get("content") or "", calls

def run_subagent(role, task):
    prompts = {
        "research": "You are a research specialist. Investigate thoroughly and report findings concisely.",
        "code": "You are a coding specialist. Write correct, minimal code and verify it with run_python.",
        "plan": "You are a planning specialist. Produce a clear, ordered, actionable plan.",
    }
    sys_p = prompts.get(role, prompts["research"]) + " Use tools when helpful. Be concise."
    msgs = [{"role": "system", "content": sys_p}, {"role": "user", "content": task}]
    model = run_subagent.model
    for _ in range(3):
        text, calls = llm(model, msgs, TOOLS_SPEC)
        if not calls:
            return text
        msgs.append({"role": "assistant", "content": text, "tool_calls": [
            {"id": f"c{i}", "type": "function", "function": {"name": c["name"], "arguments": json.dumps(c["args"])}}
            for i, c in enumerate(calls)]})
        for i, c in enumerate(calls):
            msgs.append({"role": "tool", "tool_call_id": f"c{i}", "content": run_tool(c["name"], c["args"], allow_delegate=False)})
    return text or "(sub-agent finished)"
run_subagent.model = "groq:llama-3.3-70b-versatile"

# ===========================================================================
#  Endpoints
# ===========================================================================
def web_block(text):
    try:
        out = []
        with DDGS() as d:
            for r in d.text(text, max_results=5):
                out.append({"title": r.get("title",""), "snippet": r.get("body",""), "url": r.get("href","")})
        return out
    except Exception:
        return []

@app.route("/agent", methods=["POST"])
def agent():
    data = request.get_json(force=True)
    model = data.get("model", "groq:llama-3.3-70b-versatile")
    user_msgs = data.get("messages", [])
    system = data.get("system", "") + ("\n\nYou are an autonomous agent. Break the goal into steps, "
        "call tools to gather facts or act, then give a final answer. Don't guess when a tool can confirm.")
    run_subagent.model = model
    msgs = [{"role": "system", "content": system}] + user_msgs

    def gen():
        for step in range(MAX_STEPS):
            try:
                text, calls = llm(model, msgs, TOOLS_SPEC)
            except Exception as e:
                yield json.dumps({"type": "error", "text": str(e)}) + "\n"; return
            if not calls:
                yield json.dumps({"type": "token", "text": text}) + "\n"
                yield json.dumps({"type": "done"}) + "\n"; return
            msgs.append({"role": "assistant", "content": text, "tool_calls": [
                {"id": f"c{i}", "type": "function", "function": {"name": c["name"], "arguments": json.dumps(c["args"])}}
                for i, c in enumerate(calls)]})
            for i, c in enumerate(calls):
                label = c["name"] + "(" + ", ".join(f"{k}={str(v)[:40]}" for k, v in c["args"].items()) + ")"
                yield json.dumps({"type": "step", "info": label}) + "\n"
                result = run_tool(c["name"], c["args"])
                msgs.append({"role": "tool", "tool_call_id": f"c{i}", "content": result[:4000]})
        yield json.dumps({"type": "token", "text": "I reached my step limit. Here's where I got to — ask me to continue."}) + "\n"
        yield json.dumps({"type": "done"}) + "\n"
    return Response(gen(), mimetype="application/x-ndjson")

# ---- plain streaming chat (fast path, from v2) ----
def needs_search(t):
    """Only reach for the web when the answer genuinely needs *live* info.
    The model already knows meanings, facts, and how to chat — don't waste a search on those."""
    import re
    t = t.lower().strip()
    clean = re.sub(r"[^a-z' ]", "", t).strip()
    filler = {"i don't know","i dont know","idk","no idea","not sure","i'm not sure","im not sure",
              "dunno","don't know","dont know","i don't","i dont","ok","okay","yes","no","yeah","nope",
              "thanks","thank you","never mind","nevermind","cool","nice","right","sure","maybe","hmm","hi","hey"}
    if clean in filler:
        return False
    if len(clean) < 8:                       # short conversational turns aren't searches
        return False
    # plain definition / meaning / spelling → the model knows it; don't search
    if re.search(r"^(what does|what's the meaning|whats the meaning|meaning of|define |what is the meaning|how do you spell|how to spell)", t) \
       and not re.search(r"\b(latest|today|current|recent|news|price|release|20(2[6-9]|3\d))\b", t):
        return False
    recency = [r"\blatest\b", r"\btoday\b", r"\btonight\b", r"\bcurrent(ly)?\b", r"\brecent(ly)?\b",
               r"\bnews\b", r"\bweather\b", r"\bprices?\b", r"\bstock\b", r"\bscores?\b", r"\bfixtures?\b",
               r"\bthis (week|month|year)\b", r"\bright now\b", r"\bwho won\b", r"\bwho is the current\b",
               r"\bdeadlines?\b", r"\bsearch (for|the|online|up)\b", r"\blook(ed)? (it )?up\b", r"\bgoogle it\b",
               r"\bhappening\b", r"\b20(2[6-9]|3\d)\b", r"\breleased?\b", r"\bupdates?\b"]
    return any(re.search(p, t) for p in recency)

def stream_ollama(model, msgs):
    with requests.post(OLLAMA + "/api/chat", json={"model": model, "messages": msgs, "stream": True}, stream=True, timeout=300) as r:
        for line in r.iter_lines():
            if not line: continue
            try: o = json.loads(line.decode())
            except Exception: continue
            p = o.get("message", {}).get("content", "")
            if p: yield p

def stream_compat(base, key, model_id, msgs):
    with requests.post(base + "/chat/completions",
                       headers={"Authorization": "Bearer " + key, "Content-Type": "application/json"},
                       json={"model": model_id, "messages": msgs, "stream": True}, stream=True, timeout=180) as r:
        if r.status_code != 200:
            raise RuntimeError(_api_error(r))
        for line in r.iter_lines():
            if not line: continue
            s = line.decode()
            if s.startswith("data: "): s = s[6:]
            if s.strip() == "[DONE]": break
            try:
                o = json.loads(s)
            except Exception:
                continue
            if isinstance(o, dict) and o.get("error"):
                raise RuntimeError(o["error"].get("message") if isinstance(o["error"], dict) else str(o["error"]))
            try:
                yield o["choices"][0]["delta"].get("content", "") or ""
            except Exception:
                continue

def stream_anthropic(key, model_id, msgs):
    sys_txt = " ".join(m["content"] for m in msgs if m["role"] == "system")
    conv = [{"role": m["role"], "content": m["content"]} for m in msgs if m["role"] in ("user", "assistant")]
    with requests.post("https://api.anthropic.com/v1/messages",
                       headers={"x-api-key": key, "anthropic-version": "2023-06-01", "content-type": "application/json"},
                       json={"model": model_id, "system": sys_txt, "messages": conv, "max_tokens": 1024, "stream": True},
                       stream=True, timeout=180) as r:
        if r.status_code != 200:
            raise RuntimeError(_api_error(r))
        for line in r.iter_lines():
            if not line: continue
            s = line.decode()
            if s.startswith("data: "):
                try:
                    o = json.loads(s[6:])
                    if o.get("type") == "content_block_delta":
                        yield o["delta"].get("text", "") or ""
                except Exception: continue

def _stream_one(model, msgs):
    prov, mid = split_model(model)
    if prov == "ollama":
        yield from stream_ollama(mid, msgs)
    elif prov == "claude":
        yield from stream_anthropic(ANTHROPIC_KEY, mid, msgs)
    else:
        base, keyfn = COMPAT[prov]
        yield from stream_compat(base, keyfn(), mid, msgs)

@app.route("/chat", methods=["POST"])
def chat():
    data = request.get_json(force=True)
    model = data.get("model", "phi3:mini"); messages = data.get("messages", [])
    system = data.get("system", ""); force = data.get("search", False)
    ut = next((m["content"] for m in reversed(messages) if m.get("role") == "user"), "")
    used = retrieve(ut) if ut else ""
    if used and "No documents" not in used and "Nothing relevant" not in used:
        system += "\n\nFrom Bharadwaj's documents:\n" + used
    # Markdown memory vault — keyword recall, no embedder needed.
    if ut:
        vblock, _vh = vault_search(ut)
        if vblock:
            system += "\n\nFrom Bharadwaj's memory vault (saved notes & past reports):\n" + vblock
        sk = match_skill(ut)
        if sk:
            system += "\n\nActive skill — %s (follow these for this reply):\n%s" % (sk["name"], sk["body"][:1600])
    sources = []
    if ut and (force or needs_search(ut)):
        sources = web_block(ut)
        if sources:
            system += "\n\nLive web results:\n" + "\n".join(f"[{i+1}] {s['title']} — {s['snippet']} ({s['url']})" for i, s in enumerate(sources))
    msgs = [{"role": "system", "content": system}] + messages
    def gen():
        if sources: yield json.dumps({"type": "sources", "sources": sources}) + "\n"
        chain = build_chain(model)
        last_err = "no brain available"
        for idx, mdl in enumerate(chain):
            emitted = False
            try:
                for p in _stream_one(mdl, msgs):
                    if p:
                        emitted = True
                        yield json.dumps({"type": "token", "text": p}) + "\n"
                if emitted:
                    yield json.dumps({"type": "done"}) + "\n"
                    return
                last_err = "empty response from " + mdl
            except Exception as e:
                last_err = str(e)
            # nothing usable from this brain — fall over to the next free one
            if idx + 1 < len(chain):
                yield json.dumps({"type": "switch", "to": chain[idx + 1], "why": last_err}) + "\n"
        yield json.dumps({"type": "error", "text": last_err}) + "\n"
    return Response(gen(), mimetype="application/x-ndjson")

@app.route("/ingest", methods=["POST"])
def ingest():
    folder = (request.get_json(force=True).get("path", "") or "").strip().strip('"')
    if not folder or not os.path.isdir(folder):
        return {"error": "Folder not found: " + folder}, 400
    col = get_collection()
    if col is None:
        return {"error": "Run: pip install chromadb sentence-transformers pypdf python-docx"}, 500
    exts = ("pdf","docx","txt","md","markdown","py","js","ts","json","csv","html","htm","css","xml","yaml","yml","java","sql")
    files = []
    for e in exts:
        files += glob.glob(os.path.join(folder, "**", "*." + e), recursive=True)
    files = [f for f in files if "node_modules" not in f and "tony_db" not in f][:500]
    nf = nc = 0
    for path in files:
        text = extract_text(path)
        if not text.strip(): continue
        ch = chunk(text)
        if not ch: continue
        base = os.path.basename(path)
        try:
            col.upsert(documents=ch, ids=[f"{base}_{i}_{abs(hash(path))%99999}" for i in range(len(ch))],
                       metadatas=[{"path": path, "name": base} for _ in ch])
            nf += 1; nc += len(ch)
        except Exception as e:
            print("  [upsert]", base, e)
    return {"files": nf, "chunks": nc, "total_in_db": col.count()}

@app.route("/skills", methods=["GET"])
def skills_route():
    return {"skills": [{"name": s["name"], "triggers": s["triggers"]} for s in SKILLS],
            "dir": SKILLS_DIR}

@app.route("/vault", methods=["GET"])
def vault_route():
    q = (request.args.get("q", "") or "").strip()
    if q:
        _, hits = vault_search(q, max_hits=8)
        return {"query": q, "hits": hits}
    return {"notes": vault_list(), "dir": VAULT_DIR}

@app.route("/vault/save", methods=["POST"])
def vault_save_route():
    d = request.get_json(force=True)
    content = (d.get("content", "") or "").strip()
    if not content:
        return {"error": "empty content"}, 400
    r = vault_save((d.get("title", "") or "Note").strip(), content, kind=d.get("kind", "note"))
    return {"ok": True, "name": r["name"], "slug": r["slug"]}

@app.route("/speak", methods=["POST"])
def speak_route():
    data = request.get_json(force=True)
    text = (data.get("text", "") or "")[:1800]
    voice = data.get("voice", "en-GB-RyanNeural")
    if not text.strip():
        return {"error": "empty"}, 400
    # OpenAI TTS (smoother, assistant-style) — voice like "openai:nova". Needs OPENAI_API_KEY.
    if voice.startswith("openai:"):
        if not OPENAI_KEY:
            return {"error": "OpenAI TTS needs OPENAI_API_KEY."}, 400
        try:
            r = requests.post("https://api.openai.com/v1/audio/speech",
                              headers={"Authorization": "Bearer " + OPENAI_KEY},
                              json={"model": "tts-1", "voice": voice.split(":", 1)[1] or "nova", "input": text}, timeout=60)
            if r.status_code != 200:
                return {"error": _api_error(r)}, 500
            return Response(r.content, mimetype="audio/mpeg")
        except Exception as e:
            return {"error": str(e)}, 500
    try:
        import edge_tts, asyncio
        async def synth():
            out = b""
            async for ch in edge_tts.Communicate(text, voice).stream():
                if ch["type"] == "audio":
                    out += ch["data"]
            return out
        audio = asyncio.run(synth())
        return Response(audio, mimetype="audio/mpeg")
    except ModuleNotFoundError:
        return {"error": "edge-tts not installed. Run: pip install edge-tts"}, 500
    except Exception as e:
        return {"error": str(e)}, 500


# ===========================================================================
#  Local voice (optional, offline) — Kokoro TTS + faster-whisper STT.
#  Both are lazy-loaded on first use so startup stays fast and Tony works
#  fine without them. On a no-GPU laptop they run but are noticeably slower
#  than browser speech / Edge TTS — they're here for fully-offline use.
# ===========================================================================
_KOKORO = {"pipe": None}
_WHISPER = {"model": None}

@app.route("/capabilities")
def capabilities():
    import importlib.util as _iu
    has = lambda m: _iu.find_spec(m) is not None
    return {"local_tts": has("kokoro"), "local_stt": has("faster_whisper"),
            "edge_tts": has("edge_tts")}

@app.route("/speak_local", methods=["POST"])
def speak_local():
    data = request.get_json(force=True)
    text = (data.get("text", "") or "")[:1800]
    voice = data.get("voice", "bf_emma")     # British female by default
    if not text.strip():
        return {"error": "empty"}, 400
    try:
        if _KOKORO["pipe"] is None:
            from kokoro import KPipeline
            _KOKORO["pipe"] = KPipeline(lang_code="b")    # 'b' = British English
        import numpy as np, io, wave
        chunks = [a for _, _, a in _KOKORO["pipe"](text, voice=voice)]
        if not chunks:
            return {"error": "Kokoro produced no audio."}, 500
        audio = np.concatenate(chunks)
        pcm = (np.clip(audio, -1, 1) * 32767).astype("<i2").tobytes()
        buf = io.BytesIO()
        with wave.open(buf, "wb") as w:
            w.setnchannels(1); w.setsampwidth(2); w.setframerate(24000); w.writeframes(pcm)
        return Response(buf.getvalue(), mimetype="audio/wav")
    except ModuleNotFoundError:
        return {"error": "Kokoro not installed. Run: pip install kokoro soundfile"}, 500
    except Exception as e:
        return {"error": "Kokoro TTS failed: " + str(e)}, 500

@app.route("/stt", methods=["POST"])
def stt_route():
    f = request.files.get("audio")
    if not f:
        return {"error": "no audio uploaded (form field 'audio')"}, 400
    import tempfile
    try:
        if _WHISPER["model"] is None:
            from faster_whisper import WhisperModel
            size = os.environ.get("TONY_WHISPER", "base")    # tiny|base|small
            _WHISPER["model"] = WhisperModel(size, device="cpu", compute_type="int8")
        suffix = os.path.splitext(f.filename or "audio.webm")[1] or ".webm"
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        f.save(tmp.name); tmp.close()
        segments, _info = _WHISPER["model"].transcribe(tmp.name, language="en")
        text = " ".join(s.text for s in segments).strip()
        try: os.unlink(tmp.name)
        except Exception: pass
        return {"text": text}
    except ModuleNotFoundError:
        return {"error": "faster-whisper not installed. Run: pip install faster-whisper"}, 500
    except Exception as e:
        return {"error": "STT failed: " + str(e)}, 500


@app.route("/models")
def models():
    names = []
    try:
        names = [m["name"] for m in requests.get(OLLAMA + "/api/tags", timeout=10).json().get("models", [])]
    except Exception as e:
        print("  [tags]", e)
    if GROQ_KEY:
        names += ["groq:llama-3.3-70b-versatile", "groq:llama-3.1-8b-instant"]
    if GEMINI_KEY:
        names += ["gemini:gemini-2.0-flash", "gemini:gemini-1.5-flash"]
    if OPENAI_KEY:
        names += ["openai:gpt-4o-mini", "openai:gpt-4o"]
    if ANTHROPIC_KEY:
        names += ["claude:claude-3-5-sonnet-20241022", "claude:claude-3-5-haiku-20241022"]
    if XAI_KEY:
        names += ["grok:grok-2-latest", "grok:grok-beta"]
    return {"models": [{"name": n} for n in names]}

@app.route("/")
def index(): return send_from_directory(".", "tony.html")
@app.route("/<path:p>")
def static_files(p): return send_from_directory(".", p)

# ===========================================================================
#  Background job-hunt agent — runs on a schedule, keeps a shortlist ready
# ===========================================================================
JOBS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "jobs_digest.json")
JOB_INTERVAL_HOURS = 24
JOB_QUERIES = [
    "Graduate AI Engineer jobs UK 2026",
    "Junior Machine Learning Engineer UK",
    "Graduate Data Scientist UK 2026",
    "LLM Engineer graduate UK",
    "Junior AI Engineer UK remote",
]

def pick_model():
    if GROQ_KEY:   return "groq:llama-3.3-70b-versatile"
    if GEMINI_KEY: return "gemini:gemini-2.0-flash"
    if XAI_KEY:    return "grok:grok-2-latest"
    try:
        ns = [m["name"] for m in requests.get(OLLAMA + "/api/tags", timeout=5).json().get("models", [])]
        return ns[0] if ns else None
    except Exception:
        return None

def run_job_hunt():
    import datetime
    raw, seen, uniq = [], set(), []
    for q in JOB_QUERIES:
        for r in (web_block(q) or [])[:4]:
            raw.append(r)
        time.sleep(1)
    for r in raw:
        u = r.get("url")
        if u and u not in seen:
            seen.add(u); uniq.append(r)
    digest = {"generated": datetime.datetime.now().strftime("%A %d %B %Y, %H:%M"), "summary": "", "jobs": uniq[:20]}
    model = pick_model()
    if model and uniq:
        listing = "\n".join(f"- {r['title']}: {r['snippet'][:150]} ({r['url']})" for r in uniq[:20])
        prompt = ("From these UK job results, pick the 5 strongest matches for a graduate/junior AI-ML engineer with "
                  "an MSc in AI & Robotics, a published reinforcement-learning paper, ~1 year experience, and a UK "
                  "Graduate Visa (no sponsorship needed). For each give: title, one line on why it fits, and the link. "
                  "Warm and concise.\n\n" + listing)
        try:
            txt, _ = llm(model, [{"role": "user", "content": prompt}])
            digest["summary"] = txt
        except Exception as e:
            digest["summary"] = "(summary unavailable: %s)" % e
    try:
        with open(JOBS_FILE, "w", encoding="utf-8") as f:
            json.dump(digest, f)
    except Exception as e:
        print("  [jobs write]", e)
    return digest

@app.route("/jobs")
def jobs_get():
    if os.path.exists(JOBS_FILE):
        with open(JOBS_FILE, "r", encoding="utf-8") as f:
            return Response(f.read(), mimetype="application/json")
    return {"generated": None, "summary": "No shortlist yet — tap refresh to build one.", "jobs": []}

@app.route("/jobs/refresh", methods=["POST"])
def jobs_refresh():
    return run_job_hunt()

def job_loop():
    time.sleep(10)
    while True:
        try:
            run_job_hunt(); print("  [jobs] shortlist refreshed")
        except Exception as e:
            print("  [jobs]", e)
        time.sleep(JOB_INTERVAL_HOURS * 3600)

# ===========================================================================
#  Autonomy engine — recurring routines that PREPARE only (never act)
#  Safe by design: gather, summarise, draft. Anything with consequences
#  (sending, applying, buying, posting, deleting) is left for you to approve.
# ===========================================================================
FEED_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "autonomy_feed.json")
SAFE_ROUTINE = ("You are running automatically in the background with nobody watching. "
                "You may ONLY gather, summarise and prepare information as text. You must NEVER perform any "
                "action with consequences — no sending, applying, posting, buying, deleting or changing settings. "
                "If a task would need that, prepare a draft and note that it needs Sunny's approval. UK English, warm, brief.")

def load_feed():
    try:
        with open(FEED_FILE, "r", encoding="utf-8") as f: return json.load(f)
    except Exception:
        return []

def push_feed(kind, title, body):
    import datetime
    feed = load_feed()
    feed.insert(0, {"id": str(int(time.time())), "kind": kind, "title": title, "body": body,
                    "time": datetime.datetime.now().strftime("%a %d %b · %H:%M")})
    feed = feed[:50]
    try:
        with open(FEED_FILE, "w", encoding="utf-8") as f: json.dump(feed, f)
    except Exception as e:
        print("  [feed]", e)
    # Archive substantial reports to the markdown vault so they become memory.
    try:
        if kind in ("brief", "jobs", "careeros", "review") and len(str(body)) > 40:
            vault_save("%s — %s" % (title, datetime.datetime.now().strftime("%d %b %Y")),
                       str(body), kind="report", unique=True)
    except Exception as e:
        print("  [vault archive]", e)
    return body

def routine_jobs():
    d = run_job_hunt()
    return push_feed("jobs", "Job shortlist refreshed", d.get("summary") or "Shortlist updated.")

def careeros_updates(hours=24):
    """List files changed in the CareerOS project in the last N hours."""
    if not os.path.isdir(CAREEROS_DIR):
        return ""
    cutoff = time.time() - hours * 3600; changed = []
    for root, _, files in os.walk(CAREEROS_DIR):
        if any(s in root for s in (".git", "__pycache__", ".venv", "venv", "node_modules")):
            continue
        for f in files:
            p = os.path.join(root, f)
            try:
                if os.path.getmtime(p) >= cutoff:
                    changed.append(os.path.relpath(p, CAREEROS_DIR))
            except Exception:
                pass
    if not changed:
        return ""
    changed = changed[:12]
    return "CareerOS — changed in the last %dh:\n• " % hours + "\n• ".join(changed)

def routine_careeros():
    up = careeros_updates(24)
    return push_feed("careeros", "CareerOS update", up or "No CareerOS file changes in the last 24h.")

def routine_morning_brief():
    import datetime
    today = datetime.datetime.now().strftime("%A %d %B %Y")
    shortlist = ""
    try:
        if os.path.exists(JOBS_FILE):
            jd = json.load(open(JOBS_FILE, encoding="utf-8"))
            shortlist = jd.get("summary", "") or ""
    except Exception:
        pass
    mail = ""
    try:
        if _g_token():
            mail = email_summary_text()
    except Exception:
        pass
    careeros = ""
    try:
        careeros = careeros_updates(24)
    except Exception:
        pass
    model = pick_model()
    body = shortlist or "(no shortlist yet)"
    if model:
        prompt = (SAFE_ROUTINE + "\n\nWrite Sunny's morning briefing for " + today + ". "
                  "One short warm line, then 2–4 bullets: any job emails that need attention, today's strongest job pick(s) "
                  "from the shortlist, what changed in his CareerOS project, and one encouraging nudge. Keep it tight.\n\n"
                  + ("Job emails:\n" + mail + "\n\n" if mail else "")
                  + (careeros + "\n\n" if careeros else "")
                  + "Shortlist:\n" + (shortlist or "(none yet)"))
        try:
            txt, _ = llm(model, [{"role": "user", "content": prompt}]); body = txt or body
        except Exception:
            pass
    return push_feed("brief", "Morning briefing", body)

def routine_evening_review():
    model = pick_model()
    body = "Time to wind down. Anything you want me to line up for tomorrow?"
    if model:
        prompt = (SAFE_ROUTINE + "\n\nWrite a short, warm evening check-in for Sunny: one line acknowledging the day, "
                  "then gently suggest one small thing to tee up for tomorrow's job search. Two or three sentences, no lists.")
        try:
            txt, _ = llm(model, [{"role": "user", "content": prompt}]); body = txt or body
        except Exception:
            pass
    return push_feed("review", "Evening check-in", body)

ROUTINES = [
    {"id": "jobs_am",        "label": "Morning job scan",  "at": "07:30", "fn": routine_jobs},
    {"id": "morning_brief",  "label": "Morning briefing",  "at": "08:00", "fn": routine_morning_brief},
    {"id": "careeros",       "label": "CareerOS check",    "at": "12:30", "fn": routine_careeros},
    {"id": "evening_review", "label": "Evening check-in",  "at": "18:00", "fn": routine_evening_review},
    {"id": "jobs_mid",       "label": "Midnight job scan", "at": "00:00", "fn": routine_jobs},
]
_last_run = {}   # routine id -> "YYYY-MM-DD"

def autonomy_loop():
    import datetime
    time.sleep(8)
    try:
        run_job_hunt(); print("  [auto] initial shortlist ready")
    except Exception as e:
        print("  [auto] jobs", e)
    while True:
        now = datetime.datetime.now(); hhmm = now.strftime("%H:%M"); day = now.strftime("%Y-%m-%d")
        for r in ROUTINES:
            if r["at"] == hhmm and _last_run.get(r["id"]) != day:
                _last_run[r["id"]] = day
                try:
                    r["fn"](); print("  [auto]", r["id"], "ran")
                except Exception as e:
                    print("  [auto]", r["id"], e)
        time.sleep(30)

@app.route("/feed")
def feed_get():
    return Response(json.dumps(load_feed()), mimetype="application/json")

@app.route("/feed/run", methods=["POST"])
def feed_run():
    rid = (request.get_json(force=True) or {}).get("id", "")
    for r in ROUTINES:
        if r["id"] == rid:
            try:
                return {"ok": True, "body": r["fn"]()}
            except Exception as e:
                return {"ok": False, "error": str(e)}, 500
    return {"ok": False, "error": "unknown routine"}, 404

@app.route("/routines")
def routines_get():
    return {"routines": [{"id": r["id"], "label": r["label"], "at": r["at"], "last": _last_run.get(r["id"])}
                         for r in ROUTINES]}

@app.route("/define")
def define_ep():
    return {"result": define_word(request.args.get("w", ""))}

# ===========================================================================
#  File maker — turn content into txt / md / docx / pdf, saved on the laptop
# ===========================================================================
FILES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "tony_files")
os.makedirs(FILES_DIR, exist_ok=True)

def _safe_name(n):
    import re
    return (re.sub(r"[^A-Za-z0-9._ -]", "_", n).strip().replace(" ", "_")[:80]) or "tony_file"

def _write_pdf(text, path):
    try:
        from fpdf import FPDF
        pdf = FPDF(); pdf.add_page(); pdf.set_auto_page_break(True, 15); pdf.set_font("Helvetica", size=11)
        for line in text.split("\n"):
            safe = line.encode("latin-1", "replace").decode("latin-1")
            pdf.multi_cell(0, 6, safe if safe.strip() else " ")
        pdf.output(path); return True
    except Exception:
        pass
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import cm
        c = canvas.Canvas(path, pagesize=A4); w, h = A4; y = h - 2 * cm; c.setFont("Helvetica", 11)
        for line in text.split("\n"):
            for chunk in ([line[i:i+95] for i in range(0, len(line), 95)] or [""]):
                if y < 2 * cm: c.showPage(); c.setFont("Helvetica", 11); y = h - 2 * cm
                c.drawString(2 * cm, y, chunk); y -= 14
        c.save(); return True
    except Exception:
        return False

def save_document(content, filename="tony_note", fmt="txt"):
    content = content or ""
    fmt = (fmt or "txt").lower().lstrip(".")
    stem = _safe_name(os.path.splitext(filename or "tony_note")[0])
    try:
        if fmt in ("txt", "text"):
            path = os.path.join(FILES_DIR, stem + ".txt"); open(path, "w", encoding="utf-8").write(content)
        elif fmt in ("md", "markdown"):
            path = os.path.join(FILES_DIR, stem + ".md"); open(path, "w", encoding="utf-8").write(content)
        elif fmt in ("docx", "doc", "word"):
            import docx
            d = docx.Document()
            for line in content.split("\n"):
                s = line.rstrip()
                if   s.startswith("### "): d.add_heading(s[4:], 3)
                elif s.startswith("## "):  d.add_heading(s[3:], 2)
                elif s.startswith("# "):   d.add_heading(s[2:], 1)
                else: d.add_paragraph(s)
            path = os.path.join(FILES_DIR, stem + ".docx"); d.save(path)
        elif fmt == "pdf":
            path = os.path.join(FILES_DIR, stem + ".pdf")
            if not _write_pdf(content, path):
                path = os.path.join(FILES_DIR, stem + ".txt"); open(path, "w", encoding="utf-8").write(content)
                return {"ok": True, "name": os.path.basename(path), "url": "/files/" + os.path.basename(path),
                        "note": "No PDF library found — saved as .txt instead. To enable PDFs: python -m pip install fpdf2"}
        else:
            path = os.path.join(FILES_DIR, stem + "." + _safe_name(fmt)); open(path, "w", encoding="utf-8").write(content)
        return {"ok": True, "name": os.path.basename(path), "url": "/files/" + os.path.basename(path)}
    except Exception as e:
        return {"ok": False, "error": str(e)}

@app.route("/save", methods=["POST"])
def save_ep():
    d = request.get_json(force=True) or {}
    return save_document(d.get("content", ""), d.get("filename", "tony_note"), d.get("format", "txt"))

@app.route("/files/<path:name>")
def files_ep(name):
    return send_from_directory(FILES_DIR, name, as_attachment=True)

# ===========================================================================
#  Gmail (read-only) — surface recruiter mail & job updates. Never sends.
#  Setup: console.cloud.google.com → new project → enable Gmail API →
#  OAuth consent (External, add yourself as a test user) → create OAuth
#  client (Web app) with redirect http://127.0.0.1:8000/google/callback →
#  set GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET → open /google/login once.
# ===========================================================================
def _g_save(tok):
    if "refresh_token" not in tok:
        try: tok["refresh_token"] = json.load(open(GOOGLE_TOK)).get("refresh_token", "")
        except Exception: pass
    tok["expires_at"] = time.time() + int(tok.get("expires_in", 3600)) - 60
    try:
        with open(GOOGLE_TOK, "w", encoding="utf-8") as f: json.dump(tok, f)
    except Exception as e: print("  [gmail]", e)
    return tok

def _g_token():
    if not (GOOGLE_ID and GOOGLE_SECRET): return None
    try: tok = json.load(open(GOOGLE_TOK))
    except Exception: return None
    if time.time() < tok.get("expires_at", 0): return tok.get("access_token")
    rt = tok.get("refresh_token")
    if not rt: return None
    r = requests.post("https://oauth2.googleapis.com/token",
                      data={"grant_type": "refresh_token", "refresh_token": rt,
                            "client_id": GOOGLE_ID, "client_secret": GOOGLE_SECRET}, timeout=15)
    if r.status_code != 200: return None
    return _g_save(r.json()).get("access_token")

def gmail_recent(query="", maxn=8):
    t = _g_token()
    if not t: return None
    h = {"Authorization": "Bearer " + t}
    q = query or 'newer_than:7d (recruiter OR interview OR "your application" OR application OR job OR opportunity OR hiring OR LinkedIn OR Indeed)'
    try:
        r = requests.get("https://gmail.googleapis.com/gmail/v1/users/me/messages",
                         headers=h, params={"q": q, "maxResults": maxn}, timeout=15).json()
        out = []
        for m in r.get("messages", [])[:maxn]:
            md = requests.get("https://gmail.googleapis.com/gmail/v1/users/me/messages/" + m["id"],
                              headers=h, params={"format": "metadata", "metadataHeaders": ["From", "Subject", "Date"]}, timeout=15).json()
            hs = {x["name"]: x["value"] for x in md.get("payload", {}).get("headers", [])}
            out.append({"from": hs.get("From", ""), "subject": hs.get("Subject", ""),
                        "date": hs.get("Date", ""), "snippet": md.get("snippet", "")})
        return out
    except Exception as e:
        return [{"from": "", "subject": "(error reading Gmail)", "snippet": str(e), "date": ""}]

def email_summary_text(query=""):
    import re
    items = gmail_recent(query)
    if items is None:
        return "Gmail isn't connected yet — open http://127.0.0.1:8000/google/login once."
    if not items:
        return "No job-related emails in the last 7 days. Nothing needs you right now."
    lines = ["Recent job-related email:"]
    for it in items:
        frm = re.sub(r"<.*?>", "", it.get("from", "")).strip().strip('"') or "unknown"
        lines.append("• " + (it.get("subject") or "(no subject)") + " — " + frm)
    return "\n".join(lines)

@app.route("/google/status")
def g_status():
    return {"configured": bool(GOOGLE_ID and GOOGLE_SECRET), "connected": bool(_g_token())}

@app.route("/google/login")
def g_login():
    if not (GOOGLE_ID and GOOGLE_SECRET):
        return Response("Gmail isn't set up yet. Create OAuth credentials in Google Cloud Console, set GOOGLE_CLIENT_ID and "
                        "GOOGLE_CLIENT_SECRET, add redirect http://127.0.0.1:8000/google/callback, then restart Tony.",
                        mimetype="text/plain")
    from urllib.parse import urlencode
    url = "https://accounts.google.com/o/oauth2/v2/auth?" + urlencode(
        {"client_id": GOOGLE_ID, "redirect_uri": GOOGLE_REDIRECT, "response_type": "code",
         "scope": GOOGLE_SCOPES, "access_type": "offline", "prompt": "consent"})
    return Response("", status=302, headers={"Location": url})

@app.route("/google/callback")
def g_callback():
    code = request.args.get("code", "")
    if not code:
        return Response("Google sign-in cancelled. " + request.args.get("error", ""), mimetype="text/plain")
    r = requests.post("https://oauth2.googleapis.com/token",
                      data={"grant_type": "authorization_code", "code": code, "redirect_uri": GOOGLE_REDIRECT,
                            "client_id": GOOGLE_ID, "client_secret": GOOGLE_SECRET}, timeout=15)
    if r.status_code != 200:
        return Response("Couldn't connect Gmail: " + _api_error(r), mimetype="text/plain")
    _g_save(r.json())
    return Response("<h2>Gmail connected \u2713</h2><p>Close this tab — Tony will fold your job emails into the morning briefing.</p>",
                    mimetype="text/html")

@app.route("/email")
def email_ep():
    return {"result": email_summary_text(request.args.get("q", ""))}

def careeros_report():
    if not os.path.isdir(CAREEROS_DIR):
        return "CareerOS folder not found at " + CAREEROS_DIR + ". Set TONY_CAREEROS to its path and restart."
    total = 0; newest = None; newest_t = 0
    for root, _, files in os.walk(CAREEROS_DIR):
        if any(s in root for s in (".git", "__pycache__", ".venv", "venv", "node_modules")):
            continue
        for f in files:
            p = os.path.join(root, f)
            try:
                t = os.path.getmtime(p); total += 1
                if t > newest_t: newest_t, newest = t, os.path.relpath(p, CAREEROS_DIR)
            except Exception:
                pass
    import datetime
    head = "CareerOS project: %d files." % total
    if newest:
        head += " Most recent: %s (%s)." % (newest, datetime.datetime.fromtimestamp(newest_t).strftime("%d %b %H:%M"))
    recent = careeros_updates(24)
    return head + ("\n\n" + recent if recent else "\n\nNo files changed in the last 24h.")

@app.route("/careeros")
def careeros_ep():
    return {"result": careeros_report()}

# ===========================================================================
#  Spotify — search & playback control (needs Premium + a one-time app key)
#  Setup: create an app at https://developer.spotify.com/dashboard,
#         add redirect URI  http://127.0.0.1:8000/spotify/callback,
#         then set SPOTIFY_CLIENT_ID and SPOTIFY_CLIENT_SECRET env vars.
#  Connect once by opening  http://127.0.0.1:8000/spotify/login
# ===========================================================================
SPOTIFY_TOK = os.path.join(os.path.dirname(os.path.abspath(__file__)), "spotify_token.json")
SPOTIFY_SCOPES = "user-modify-playback-state user-read-playback-state streaming"

def _sp_save(tok):
    if "refresh_token" not in tok:                # refresh responses omit it — keep the old one
        try: tok["refresh_token"] = json.load(open(SPOTIFY_TOK)).get("refresh_token", "")
        except Exception: pass
    tok["expires_at"] = time.time() + int(tok.get("expires_in", 3600)) - 60
    try:
        with open(SPOTIFY_TOK, "w", encoding="utf-8") as f: json.dump(tok, f)
    except Exception as e: print("  [spotify]", e)
    return tok

def _sp_token():
    """Valid access token, refreshing if needed. Returns None if not connected/configured."""
    if not (SPOTIFY_ID and SPOTIFY_SECRET): return None
    try: tok = json.load(open(SPOTIFY_TOK))
    except Exception: return None
    if time.time() < tok.get("expires_at", 0): return tok.get("access_token")
    rt = tok.get("refresh_token")
    if not rt: return None
    r = requests.post("https://accounts.spotify.com/api/token",
                      data={"grant_type": "refresh_token", "refresh_token": rt},
                      auth=(SPOTIFY_ID, SPOTIFY_SECRET), timeout=15)
    if r.status_code != 200: return None
    return _sp_save(r.json()).get("access_token")

def _sp_headers():
    t = _sp_token()
    return {"Authorization": "Bearer " + t} if t else None

@app.route("/spotify/status")
def sp_status():
    return {"configured": bool(SPOTIFY_ID and SPOTIFY_SECRET), "connected": bool(_sp_token())}

@app.route("/spotify/login")
def sp_login():
    if not (SPOTIFY_ID and SPOTIFY_SECRET):
        return Response("Spotify isn't set up yet. Create an app at developer.spotify.com, add the redirect URI "
                        "http://127.0.0.1:8000/spotify/callback, then set SPOTIFY_CLIENT_ID and SPOTIFY_CLIENT_SECRET and restart Tony.",
                        mimetype="text/plain")
    from urllib.parse import urlencode
    url = "https://accounts.spotify.com/authorize?" + urlencode({
        "client_id": SPOTIFY_ID, "response_type": "code", "redirect_uri": SPOTIFY_REDIRECT, "scope": SPOTIFY_SCOPES})
    return Response("", status=302, headers={"Location": url})

@app.route("/spotify/callback")
def sp_callback():
    code = request.args.get("code", "")
    if not code:
        return Response("Spotify authorisation was cancelled. You can close this tab.", mimetype="text/plain")
    r = requests.post("https://accounts.spotify.com/api/token",
                      data={"grant_type": "authorization_code", "code": code, "redirect_uri": SPOTIFY_REDIRECT},
                      auth=(SPOTIFY_ID, SPOTIFY_SECRET), timeout=15)
    if r.status_code != 200:
        return Response("Couldn't connect Spotify: " + _api_error(r), mimetype="text/plain")
    _sp_save(r.json())
    return Response("<h2>Spotify connected ✓</h2><p>You can close this tab and tell Tony to play music.</p>",
                    mimetype="text/html")

def _sp_active_device(h):
    try:
        ds = requests.get("https://api.spotify.com/v1/me/player/devices", headers=h, timeout=10).json().get("devices", [])
    except Exception:
        ds = []
    if not ds: return None
    act = next((d for d in ds if d.get("is_active")), None)
    return (act or ds[0]).get("id")

@app.route("/spotify/play", methods=["POST"])
def sp_play():
    h = _sp_headers()
    if not h: return {"ok": False, "error": "Spotify isn't connected. Open http://127.0.0.1:8000/spotify/login"}, 400
    q = (request.get_json(force=True) or {}).get("query", "").strip()
    body, msg = None, "Resuming."
    if q:
        sr = requests.get("https://api.spotify.com/v1/search", headers=h,
                          params={"q": q, "type": "track", "limit": 1}, timeout=10).json()
        items = sr.get("tracks", {}).get("items", [])
        if not items: return {"ok": False, "error": "Couldn't find “%s” on Spotify." % q}, 404
        tr = items[0]; body = {"uris": [tr["uri"]]}
        msg = "Playing %s by %s." % (tr["name"], ", ".join(a["name"] for a in tr["artists"]))
    dev = _sp_active_device(h)
    url = "https://api.spotify.com/v1/me/player/play" + (("?device_id=" + dev) if dev else "")
    pr = requests.put(url, headers=h, json=body or {}, timeout=10)
    if pr.status_code in (403,):
        return {"ok": False, "error": "Playback needs Spotify Premium."}, 403
    if pr.status_code == 404 or not dev:
        return {"ok": False, "error": "No active Spotify device — open Spotify on your phone or desktop first, then try again."}, 404
    if pr.status_code >= 400:
        return {"ok": False, "error": _api_error(pr)}, pr.status_code
    return {"ok": True, "message": msg}

def _sp_simple(method_path, ok_msg):
    h = _sp_headers()
    if not h: return {"ok": False, "error": "Spotify isn't connected."}, 400
    verb, path = method_path
    r = (requests.put if verb == "put" else requests.post)("https://api.spotify.com/v1/me/player/" + path,
                                                            headers=h, timeout=10)
    if r.status_code == 403: return {"ok": False, "error": "Needs Spotify Premium."}, 403
    if r.status_code >= 400 and r.status_code != 404: return {"ok": False, "error": _api_error(r)}, r.status_code
    return {"ok": True, "message": ok_msg}

@app.route("/spotify/pause", methods=["POST"])
def sp_pause(): return _sp_simple(("put", "pause"), "Paused.")
@app.route("/spotify/next", methods=["POST"])
def sp_next(): return _sp_simple(("post", "next"), "Skipping ahead.")
@app.route("/spotify/prev", methods=["POST"])
def sp_prev(): return _sp_simple(("post", "previous"), "Going back.")

def sp_play_text(q=""):
    """String result for the agent's play_music tool (reuses the same auth/device logic)."""
    h = _sp_headers()
    if not h: return "Spotify isn't connected yet — open http://127.0.0.1:8000/spotify/login once."
    q = (q or "").strip(); body = None; msg = "Resuming."
    if q:
        sr = requests.get("https://api.spotify.com/v1/search", headers=h,
                          params={"q": q, "type": "track", "limit": 1}, timeout=10).json()
        items = sr.get("tracks", {}).get("items", [])
        if not items: return "Couldn't find “%s” on Spotify." % q
        tr = items[0]; body = {"uris": [tr["uri"]]}
        msg = "Playing %s by %s." % (tr["name"], ", ".join(a["name"] for a in tr["artists"]))
    dev = _sp_active_device(h)
    url = "https://api.spotify.com/v1/me/player/play" + (("?device_id=" + dev) if dev else "")
    pr = requests.put(url, headers=h, json=body or {}, timeout=10)
    if pr.status_code == 403: return "Playback needs Spotify Premium."
    if pr.status_code == 404 or not dev: return "No active Spotify device — open Spotify on a device first, then ask again."
    if pr.status_code >= 400: return "Spotify: " + _api_error(pr)
    return "▶ " + msg

if __name__ == "__main__":
    try:
        import playwright; pw = "ready"
    except Exception:
        pw = "off (pip install playwright && playwright install chromium)"
    brains = ["ollama (local)"]
    if GROQ_KEY: brains.append("groq")
    if GEMINI_KEY: brains.append("gemini")
    if OPENAI_KEY: brains.append("openai")
    if ANTHROPIC_KEY: brains.append("claude")
    if XAI_KEY: brains.append("grok")
    print("\n  TONY backend (v3 · agentic) →  http://localhost:%d" % PORT)
    print("  Brains     : " + ", ".join(brains))
    print("  Browser    : " + pw)
    if HOST == "127.0.0.1":
        print("  Security   : localhost only (other devices can't reach it) ✓")
        print("               For phone access: set TONY_LAN=1  (and ideally TONY_TOKEN=some-secret)")
    else:
        print("  Security   : EXPOSED on your network (%s) — other devices on this WiFi can reach Tony." % HOST)
        print("               Phone link: http://<your-laptop-IP>:%d   (run 'ipconfig' for the IPv4)" % PORT)
        if TOKEN:
            print("               Token ON — on each new device open: http://<laptop-IP>:%d/?token=YOUR_TOKEN once ✓" % PORT)
        else:
            print("               ⚠ No TONY_TOKEN set — anyone on this WiFi can use Tony. Set one if you're not on a trusted network.")
    print("  Autonomy   : morning scan 07:30 · briefing 08:00 · evening 18:00 · midnight scan 00:00 (prepare-only)")
    if SPOTIFY_ID and SPOTIFY_SECRET:
        print("  Spotify    : configured — connect once at http://127.0.0.1:%d/spotify/login (Premium needed to play)" % PORT)
    else:
        print("  Spotify    : off — set SPOTIFY_CLIENT_ID & SPOTIFY_CLIENT_SECRET to enable 'play music'")
    print("  Make sure Ollama is running.\n")
    def _warm():
        print("  [docs] loading embedder (first run downloads ~80MB)…")
        if get_collection() is not None: print("  [docs] embedder ready ✓")
        else: print("  [docs] embedder unavailable — run: python -m pip install sentence-transformers chromadb")
    threading.Thread(target=_warm, daemon=True).start()
    threading.Thread(target=autonomy_loop, daemon=True).start()
    app.run(host=HOST, port=PORT, threaded=True)